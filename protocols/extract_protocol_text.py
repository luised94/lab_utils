# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "pymupdf",
#   "python-docx",
#   "pytesseract",
#   "Pillow",
# ]
# ///

import os
import sys
import shutil
import time
from pathlib import Path

import fitz
import docx
import pytesseract
from PIL import Image

# =============================================================================
# CONFIGURATION
# =============================================================================

INPUT_DIR              = Path.home() / "lab" / "protocols"
OUTPUT_DIR_SUFFIX      = "_extracted"
SCANNED_PAGE_THRESHOLD = 50
TESSERACT_LANG         = "eng"
SUPPORTED_EXTENSIONS   = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".docx"}
WARN_EXTENSIONS        = {".doc"}

# =============================================================================
# DERIVED VALUES
# =============================================================================

output_dir = INPUT_DIR.parent / (INPUT_DIR.name + OUTPUT_DIR_SUFFIX)

# =============================================================================
# PREFLIGHT
# =============================================================================

if shutil.which("tesseract") is None:
    print("ERROR: tesseract-ocr not found on PATH.")
    print("       Install with: sudo apt install -y tesseract-ocr")
    sys.exit(1)

if not INPUT_DIR.exists():
    print(f"ERROR: Input directory not found: {INPUT_DIR}")
    sys.exit(1)

if not INPUT_DIR.is_dir():
    print(f"ERROR: {INPUT_DIR} exists but is not a directory.")
    sys.exit(1)

if not os.access(INPUT_DIR, os.R_OK):
    print(f"ERROR: No read permission on {INPUT_DIR}")
    sys.exit(1)

if not os.access(output_dir.parent, os.W_OK):
    print(f"ERROR: No write permission in {output_dir.parent}")
    print(f"       Cannot create output directory.")
    sys.exit(1)

output_dir.mkdir(exist_ok=True)

print(f"INFO: Input  : {INPUT_DIR}")
print(f"INFO: Output : {output_dir}")
# =============================================================================
# METRICS
# =============================================================================

start_time        = time.monotonic()

total_seen        = 0
total_attempted   = 0
total_processed   = 0
total_errored     = 0
total_skipped     = 0
total_warned      = 0
total_unsupported = 0

count_pdf         = 0
count_image       = 0
count_docx        = 0

total_chars       = 0

# =============================================================================
# MAIN LOOP
# =============================================================================

all_files = sorted(INPUT_DIR.iterdir())

for file in all_files:

    if not file.is_file():
        continue

    total_seen += 1
    ext = file.suffix.lower()

    if ext in WARN_EXTENSIONS:
        print(f"WARNING: {file.name} is .doc - not supported. Convert to .docx.")
        total_warned += 1
        continue

    if ext not in SUPPORTED_EXTENSIONS:
        print(f"INFO   : {file.name} - unsupported type, skipping.")
        total_unsupported += 1
        continue

    out_path = output_dir / (file.stem + ".txt")

    if out_path.exists():
        print(f"WARNING: {out_path.name} already exists - skipping {file.name}.")
        total_skipped += 1
        continue

    total_attempted += 1


    try:
        text = None

        if ext == ".pdf":
            # extraction: commits 5-6
            pass

        elif ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
            # extraction: commit 4
            pass

        elif ext == ".docx":
            doc  = docx.Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += "\t".join(cell.text.strip() for cell in row.cells) + "\n"
            count_docx += 1

        if text is not None:
            out_path.write_text(text, encoding="utf-8")
            total_chars += len(text)
            total_processed += 1
            print(f"OK     : {file.name}  {out_path.name}")

    except Exception as e:
        total_errored += 1
        print(f"ERROR  : {file.name} failed - {e}")
# =============================================================================
# SUMMARY
# =============================================================================

elapsed = time.monotonic() - start_time

print("-" * 45)
print(f"Extraction complete ({elapsed:.1f}s)")
print("-" * 45)
print(f"  Files seen     : {total_seen}")
print(f"  Extracted      : {total_processed}  (pdf: {count_pdf}  image: {count_image}  docx: {count_docx})")
print(f"  Errored        : {total_errored}")
print(f"  Skipped        : {total_skipped}  (output existed)")
print(f"  .doc warned    : {total_warned}")
print(f"  Unsupported    : {total_unsupported}")
print("-" * 45)
print(f"  Total chars    : {total_chars:,}")
print(f"  Output dir     : {output_dir}")
print("-" * 45)

if total_errored > 0:
    print("HINT: errored files were not written - rerun to retry them.")
